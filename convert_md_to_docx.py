# -*- coding: utf-8 -*-
"""
MarkdownファイルをDOCX形式に変換するスクリプト
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """段落にハイパーリンクを追加"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # ハイパーリンクのスタイル
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """MarkdownファイルをDOCX形式に変換"""

    # Wordドキュメントを作成
    doc = Document()

    # フォント設定（日本語対応）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Gothic'
    font.size = Pt(11)

    # Markdownファイルを読み込み
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # コードブロックの処理
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph(line, style='Normal')
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(10)
            i += 1
            continue

        # 空行
        if not line.strip():
            if in_table:
                # テーブル終了
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
            doc.add_paragraph()
            i += 1
            continue

        # 見出し
        if line.startswith('#'):
            if in_table:
                process_table(doc, table_lines)
                table_lines = []
                in_table = False

            heading_level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('#').strip()

            if heading_level == 1:
                p = doc.add_heading(heading_text, level=1)
            elif heading_level == 2:
                p = doc.add_heading(heading_text, level=2)
            elif heading_level == 3:
                p = doc.add_heading(heading_text, level=3)
            else:
                p = doc.add_heading(heading_text, level=4)
            i += 1
            continue

        # テーブル
        if '|' in line and not line.startswith('>'):
            in_table = True
            table_lines.append(line)
            i += 1
            continue

        # 引用（ブロッククォート）
        if line.startswith('>'):
            if in_table:
                process_table(doc, table_lines)
                table_lines = []
                in_table = False

            quote_text = line.lstrip('>').strip()
            p = doc.add_paragraph(quote_text)
            p.style.font.italic = True
            p.style.font.color.rgb = RGBColor(128, 128, 128)
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # リスト
        if re.match(r'^[\-\*\+]\s', line) or re.match(r'^\d+\.\s', line):
            if in_table:
                process_table(doc, table_lines)
                table_lines = []
                in_table = False

            # リストアイテムのテキストを取得
            list_text = re.sub(r'^[\-\*\+\d\.]\s+', '', line).strip()

            # ネストレベルを検出
            indent_level = (len(line) - len(line.lstrip())) // 2

            p = doc.add_paragraph(list_text, style='List Bullet' if re.match(r'^[\-\*\+]\s', line.lstrip()) else 'List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            i += 1
            continue

        # 水平線
        if line.strip() in ['---', '***', '___']:
            if in_table:
                process_table(doc, table_lines)
                table_lines = []
                in_table = False

            p = doc.add_paragraph()
            p.add_run('_' * 50)
            i += 1
            continue

        # 通常の段落
        if in_table:
            process_table(doc, table_lines)
            table_lines = []
            in_table = False

        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        i += 1

    # 最後にテーブルが残っている場合
    if in_table and table_lines:
        process_table(doc, table_lines)

    # ドキュメントを保存
    doc.save(docx_file_path)
    print(f"DOCXファイルを保存しました: {docx_file_path}")

def process_table(doc, table_lines):
    """テーブルをドキュメントに追加"""
    if not table_lines:
        return

    # ヘッダー行と区切り行を除外
    header_line = table_lines[0] if table_lines else None
    data_lines = [line for line in table_lines[2:] if '|' in line] if len(table_lines) > 2 else []

    if not header_line:
        return

    # 列数を計算
    headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
    num_cols = len(headers)

    if num_cols == 0:
        return

    # テーブルを作成
    num_rows = len(data_lines) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # ヘッダー行
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # データ行
    for row_idx, data_line in enumerate(data_lines, start=1):
        cells_data = [cell.strip() for cell in data_line.split('|')[1:-1]]
        for col_idx, cell_data in enumerate(cells_data):
            if col_idx < num_cols:
                table.rows[row_idx].cells[col_idx].text = cell_data

def process_inline_formatting(paragraph, text):
    """インラインフォーマット（太字、斜体、コードなど）を処理"""
    # **太字** を処理
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # リンクを処理
            link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
            if re.search(link_pattern, part):
                # リンクがある場合
                last_pos = 0
                for match in re.finditer(link_pattern, part):
                    # リンク前のテキスト
                    if match.start() > last_pos:
                        paragraph.add_run(part[last_pos:match.start()])

                    # リンク
                    link_text = match.group(1)
                    link_url = match.group(2)

                    # URLが#で始まる場合は内部リンクなので通常のテキストとして扱う
                    if link_url.startswith('#'):
                        paragraph.add_run(link_text)
                    else:
                        add_hyperlink(paragraph, link_text, link_url)

                    last_pos = match.end()

                # リンク後のテキスト
                if last_pos < len(part):
                    paragraph.add_run(part[last_pos:])
            else:
                paragraph.add_run(part)

if __name__ == '__main__':
    md_file = r'd:\ts_pm_all_v2\docs\TIERA_USER_MANUAL.md'
    docx_file = r'd:\ts_pm_all_v2\docs\TIERA_USER_MANUAL.docx'

    print(f"Markdownファイル: {md_file}")
    print(f"DOCXファイル: {docx_file}")
    print()

    convert_markdown_to_docx(md_file, docx_file)
